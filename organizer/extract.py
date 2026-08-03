"""Extração de um trecho de texto do arquivo, para alimentar o LLM (Fase 3).

`pdfplumber` e `python-docx` são importados **dentro** das funções: quem chama
`extract` é sempre um processo filho efêmero, nunca o watcher (RNF-03).

Nada aqui executa o arquivo (RF-27) e nada levanta exceção: falha é dado, não
acidente — `Extracao.ok=False` com um motivo legível.

Requisitos cobertos: RF-48, RF-49, RF-50.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

#: Teto de caracteres devolvidos (RF-48).
MAX_CHARS = 500
#: Páginas de PDF lidas no máximo (RF-48).
MAX_PAGINAS = 2
#: Cascata de encoding para texto puro (RF-50).
ENCODINGS = ("utf-8", "utf-8-sig", "cp1252", "latin-1")

EXTENSOES_TEXTO_PURO = frozenset({".txt", ".md", ".csv", ".tsv", ".log"})
EXTENSAO_PDF = ".pdf"
EXTENSAO_DOCX = ".docx"

MOTIVO_OK = "extraido"
MOTIVO_PROTEGIDO = "protegido"
MOTIVO_VAZIO = "sem_texto"
MOTIVO_NAO_SUPORTADO = "nao_suportado"
MOTIVO_FALHOU = "falhou"

#: Quanto do começo e do fim do arquivo é varrido atrás do dicionário `/Encrypt`.
_JANELA_TRAILER = 8192


@dataclass(frozen=True)
class Extracao:
    """Resultado da extração. `ok=False` nunca é exceção — é dado."""

    texto: str
    ok: bool
    motivo: str

    def __bool__(self) -> bool:
        return self.ok


def _limitar(texto: str, max_chars: int) -> str:
    """Remove controles, colapsa espaços e corta no teto (RF-48)."""
    limpo = "".join(ch if ch >= " " or ch in "\n\t" else " " for ch in texto)
    return " ".join(limpo.split())[:max_chars]


# --------------------------------------------------------------------------- #
# Texto puro
# --------------------------------------------------------------------------- #


def ler_texto_puro(caminho: Path, max_chars: int = MAX_CHARS) -> Extracao:
    """Lê `.txt`/`.md`/`.csv` tentando os encodings em cascata (RF-50).

    Nunca levanta `UnicodeDecodeError`: o último recurso é `latin-1`, que aceita
    qualquer sequência de bytes.
    """
    bruto = caminho.read_bytes()[: max_chars * 8]
    for codec in ENCODINGS:
        try:
            return Extracao(_limitar(bruto.decode(codec), max_chars), True, MOTIVO_OK)
        except UnicodeDecodeError:
            continue
    return Extracao(  # pragma: no cover - latin-1 aceita qualquer byte
        _limitar(bruto.decode("latin-1", "replace"), max_chars), True, MOTIVO_OK
    )


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


def pdf_esta_protegido(caminho: Path) -> bool:
    """Detecta criptografia lendo o trailer, sem depender do parser nem de senha.

    Um PDF cifrado sempre declara `/Encrypt` no trailer. Checar isso direto é
    barato e responde antes de o `pdfminer` tropeçar num filtro que não sabe
    decifrar — o que produziria um erro genérico, não "protegido".
    """
    try:
        with open(caminho, "rb") as fh:
            inicio = fh.read(_JANELA_TRAILER)
            fh.seek(0, 2)
            tamanho = fh.tell()
            fh.seek(max(0, tamanho - _JANELA_TRAILER))
            fim = fh.read(_JANELA_TRAILER)
    except OSError:  # pragma: no cover - arquivo sumiu entre o is_file e o open
        return False
    return b"/Encrypt" in inicio or b"/Encrypt" in fim


def _erro_de_senha(exc: Exception) -> bool:
    """Segunda linha: o `pdfminer` também sinaliza senha por tipo de exceção."""
    nomes = {tipo.__name__ for tipo in type(exc).__mro__}
    if any("Password" in nome for nome in nomes):
        return True
    return "password" in str(exc).lower()


def ler_pdf(caminho: Path, max_chars: int = MAX_CHARS, max_paginas: int = MAX_PAGINAS) -> Extracao:
    """Texto das `max_paginas` primeiras páginas. PDF com senha → `protegido` (RF-49)."""
    if pdf_esta_protegido(caminho):
        return Extracao("", False, MOTIVO_PROTEGIDO)
    try:
        import pdfplumber

        partes: list[str] = []
        with pdfplumber.open(caminho) as pdf:
            for pagina in pdf.pages[:max_paginas]:
                texto = pagina.extract_text() or ""
                if texto:
                    partes.append(texto)
                if sum(len(p) for p in partes) >= max_chars:
                    break
    except Exception as exc:
        return Extracao("", False, MOTIVO_PROTEGIDO if _erro_de_senha(exc) else MOTIVO_FALHOU)

    texto = _limitar("\n".join(partes), max_chars)
    return Extracao(texto, True, MOTIVO_OK) if texto else Extracao("", False, MOTIVO_VAZIO)


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


def ler_docx(caminho: Path, max_chars: int = MAX_CHARS) -> Extracao:
    """Parágrafos e células de tabela do `.docx`, na ordem em que aparecem."""
    try:
        from docx import Document

        documento = Document(str(caminho))
        partes = [p.text for p in documento.paragraphs if p.text.strip()]
        for tabela in documento.tables:
            if sum(len(p) for p in partes) >= max_chars:
                break
            for linha in tabela.rows:
                partes.extend(c.text for c in linha.cells if c.text.strip())
    except Exception:
        return Extracao("", False, MOTIVO_FALHOU)

    texto = _limitar("\n".join(partes), max_chars)
    return Extracao(texto, True, MOTIVO_OK) if texto else Extracao("", False, MOTIVO_VAZIO)


# --------------------------------------------------------------------------- #
# Entrada
# --------------------------------------------------------------------------- #


def suportado(caminho: Path | str) -> bool:
    """Diz se sabemos extrair texto desta extensão."""
    ext = Path(caminho).suffix.lower()
    return ext in (EXTENSAO_PDF, EXTENSAO_DOCX) or ext in EXTENSOES_TEXTO_PURO


def trecho(caminho: Path | str, max_chars: int = MAX_CHARS) -> Extracao:
    """Trecho inicial do arquivo. Nunca executa nada e nunca levanta exceção."""
    alvo = Path(caminho)
    ext = alvo.suffix.lower()
    try:
        if not alvo.is_file():
            return Extracao("", False, MOTIVO_FALHOU)
        if ext == EXTENSAO_PDF:
            return ler_pdf(alvo, max_chars)
        if ext == EXTENSAO_DOCX:
            return ler_docx(alvo, max_chars)
        if ext in EXTENSOES_TEXTO_PURO:
            return ler_texto_puro(alvo, max_chars)
    except OSError:  # pragma: no cover - I/O quebrou no meio
        return Extracao("", False, MOTIVO_FALHOU)
    return Extracao("", False, MOTIVO_NAO_SUPORTADO)
